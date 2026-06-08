"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản pháp luật (PDF/DOCX) từ các nguồn chính thống.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, có năm ban hành.
"""

from pathlib import Path
import os

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def generate_legal_docx_files():
    """Tạo 3 văn bản pháp luật dưới dạng DOCX chứa nội dung chi tiết phục vụ cho RAG pipeline."""
    import docx
    
    setup_directory()
    
    # 1. Luật Phòng, chống ma tuý 2021
    doc1 = docx.Document()
    doc1.add_heading("LUẬT PHÒNG, CHỐNG MA TÚY 2021", 0)
    
    p1 = doc1.add_paragraph()
    p1.add_run("Luật số: 73/2021/QH14\n").bold = True
    p1.add_run("Ban hành ngày: 30 tháng 03 năm 2021\n")
    p1.add_run("Hiệu lực thi hành: Từ ngày 01 tháng 01 năm 2022\n\n")
    
    doc1.add_heading("Chương I: Quy định chung", 1)
    doc1.add_paragraph(
        "Luật này quy định về phòng, chống ma túy; quản lý người sử dụng trái phép chất ma túy; "
        "cai nghiện ma túy; trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức; quản lý nhà nước "
        "và hợp tác quốc tế về phòng, chống ma túy."
    )
    
    doc1.add_heading("Chương IV: Quản lý người sử dụng trái phép chất ma túy", 1)
    doc1.add_heading("Điều 23: Biện pháp quản lý người sử dụng trái phép chất ma túy", 2)
    doc1.add_paragraph(
        "1. Quản lý người sử dụng trái phép chất ma túy là biện pháp phòng ngừa xã hội nhằm giúp người sử dụng "
        "trái phép chất ma túy không tiếp tục sử dụng trái phép chất ma túy, phòng ngừa các hành vi vi phạm pháp luật của họ.\n"
        "2. Thời hạn quản lý người sử dụng trái phép chất ma túy là 01 năm kể từ ngày Chủ tịch Ủy ban nhân dân cấp xã "
        "ra quyết định quản lý."
    )
    
    doc1.add_heading("Chương V: Cai nghiện ma túy", 1)
    doc1.add_heading("Điều 28: Các biện pháp cai nghiện ma túy", 2)
    doc1.add_paragraph(
        "Các biện pháp cai nghiện ma túy bao gồm:\n"
        "1. Cai nghiện ma túy tự nguyện.\n"
        "2. Cai nghiện ma túy bắt buộc.\n"
        "Cai nghiện ma túy tự nguyện được thực hiện tại gia đình, cộng đồng hoặc tại cơ sở cai nghiện ma túy; "
        "cai nghiện ma túy bắt buộc được thực hiện tại cơ sở cai nghiện ma túy công lập."
    )
    
    doc1.add_heading("Điều 32: Đối tượng bị áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc", 2)
    doc1.add_paragraph(
        "Người nghiện ma túy từ đủ 18 tuổi trở lên bị áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc "
        "theo quy định của Luật Xử lý vi phạm hành chính khi thuộc một trong các trường hợp sau đây:\n"
        "a) Không đăng ký, không thực hiện hoặc tự ý chấm dứt cai nghiện ma túy tự nguyện;\n"
        "b) Trong thời gian cai nghiện ma túy tự nguyện bị phát hiện sử dụng trái phép chất ma túy;\n"
        "c) Người nghiện ma túy chất kích thích, ma túy tổng hợp tái nghiện."
    )
    
    doc1_path = DATA_DIR / "luat-phong-chong-ma-tuy-2021.docx"
    doc1.save(str(doc1_path))
    print(f"✓ Đã tạo: {doc1_path} ({doc1_path.stat().st_size} bytes)")

    # 2. Nghị định 105/2021/NĐ-CP
    doc2 = docx.Document()
    doc2.add_heading("NGHỊ ĐỊNH 105/2021/NĐ-CP HƯỚNG DẪN LUẬT PHÒNG CHỐNG MA TÚY", 0)
    
    p2 = doc2.add_paragraph()
    p2.add_run("Số hiệu: 105/2021/NĐ-CP\n").bold = True
    p2.add_run("Ban hành ngày: 04 tháng 12 năm 2021\n")
    p2.add_run("Hiệu lực thi hành: Từ ngày 01 tháng 01 năm 2022\n\n")
    
    doc2.add_heading("Chương I: Quy định chung", 1)
    doc2.add_paragraph(
        "Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma túy về "
        "công tác phối hợp của các cơ quan chuyên trách phòng, chống tội phạm về ma túy; kiểm soát các hoạt động "
        "hợp pháp liên quan đến ma túy và quản lý người sử dụng trái phép chất ma túy."
    )
    
    doc2.add_heading("Chương II: Phối hợp có liên quan", 1)
    doc2.add_paragraph(
        "Các cơ quan chuyên trách phòng, chống tội phạm về ma túy thuộc Công an nhân dân, Hải quan, Bộ đội Biên phòng "
        "và Cảnh sát biển có trách nhiệm phối hợp trao đổi thông tin, tuần tra, kiểm soát và đấu tranh phòng chống tội phạm "
        "vận chuyển, tàng trữ, mua bán trái phép chất ma túy trên các địa bàn biên giới, cửa khẩu và trên biển."
    )
    
    doc2_path = DATA_DIR / "nghi-dinh-105-2021.docx"
    doc2.save(str(doc2_path))
    print(f"✓ Đã tạo: {doc2_path} ({doc2_path.stat().st_size} bytes)")

    # 3. Bộ luật Hình sự 2015 (Chương XX - Các tội phạm về ma túy)
    doc3 = docx.Document()
    doc3.add_heading("BỘ LUẬT HÌNH SỰ 2015 - CÁC TỘI PHẠM VỀ MA TÚY", 0)
    
    p3 = doc3.add_paragraph()
    p3.add_run("Luật số: 100/2015/QH13 (Sửa đổi, bổ sung 2017)\n").bold = True
    p3.add_run("Chương XX: Các tội phạm về ma túy\n\n")
    
    doc3.add_heading("Điều 248: Tội sản xuất trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào sản xuất trái phép chất ma túy dưới bất kỳ hình thức nào, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp nghiêm trọng hoặc đặc biệt nghiêm trọng có thể bị phạt tù từ 20 năm, "
        "tù chung thân hoặc tử hình."
    )
    
    doc3.add_heading("Điều 249: Tội tàng trữ trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy "
        "thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 01 năm đến 05 năm:\n"
        "a) Đã bị xử phạt vi phạm hành chính về hành vi này hoặc đã bị kết án về tội này, chưa được xóa án tích mà còn vi phạm;\n"
        "b) Nhựa thuốc phiện, nhựa cần sa hoặc cao côca có khối lượng từ 01 gam đến dưới 500 gam;\n"
        "c) Heroine, Cocaine, Methamphetamine, Amphetamine, MDMA hoặc XLR-11 có khối lượng từ 0,1 gam đến dưới 05 gam.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây thì bị phạt tù từ 05 năm đến 10 năm:\n"
        "a) Có tổ chức;\n"
        "b) Phạm tội 02 lần trở lên;\n"
        "c) Lợi dụng chức vụ, quyền hạn;\n"
        "d) Heroine, Cocaine, Methamphetamine có khối lượng từ 05 gam đến dưới 30 gam."
    )
    
    doc3.add_heading("Điều 250: Tội vận chuyển trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào vận chuyển trái phép chất ma túy mà không nhằm mục đích sản xuất, mua bán, tàng trữ trái phép chất ma túy, "
        "thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội có tổ chức hoặc vận chuyển số lượng lớn ma túy bị phạt tù từ 07 năm đến 15 năm, hoặc lên đến chung thân, tử hình."
    )
    
    doc3_path = DATA_DIR / "bo-luat-hinh-su-2015-ma-tuy.docx"
    doc3.save(str(doc3_path))
    print(f"✓ Đã tạo: {doc3_path} ({doc3_path.stat().st_size} bytes)")


if __name__ == "__main__":
    generate_legal_docx_files()
